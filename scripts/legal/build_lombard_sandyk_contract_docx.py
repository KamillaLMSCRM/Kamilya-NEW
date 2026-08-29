"""Build the customer-ready Kamilya LMS contract pack from reviewed Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
LEGAL = ROOT / "docs" / "legal"
OUTPUT = LEGAL / "Kamilya_LMS_Dogovor_Lombard_Sandyk_RU.docx"
SOURCES = (
    LEGAL / "lombard-sandyk-b2b-agreement-ru.md",
    LEGAL / "lombard-sandyk-order-form-ru.md",
    LEGAL / "lombard-sandyk-dpa-ru.md",
)

FONT = "Arial"
INK = RGBColor(28, 37, 48)
NAVY = RGBColor(31, 58, 95)
MUTED = RGBColor(95, 105, 118)
LIGHT_FILL = "F2F4F7"
PLACEHOLDER_FILL = "FFF2CC"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size: float = 11, bold: bool | None = None, color=INK, italic: bool | None = None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_run(run, fill: str = PLACEHOLDER_FILL):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    run._element.get_or_add_rPr().append(shd)


def add_inline_runs(paragraph, text: str, *, size: float = 11, color=INK):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\[[^\]]+\])")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color)
        else:
            run = paragraph.add_run(token)
            set_run_font(run, size=size, color=color)
            shade_run(run)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in (
        ("Heading 1", 14, 14, 7),
        ("Heading 2", 12, 10, 5),
        ("Heading 3", 11, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)


def style_table(table):
    table.style = "Table Grid"
    col_count = len(table.columns)
    if col_count == 2:
        widths = [2880, 6480]
    else:
        base = TABLE_WIDTH_DXA // col_count
        widths = [base] * col_count
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)

    for cell in table.rows[0].cells:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT_FILL)
        cell._tc.get_or_add_tcPr().append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True, color=INK)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, size=10, color=INK)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, value, end))
    set_run_font(run, size=9, color=MUTED)


def configure_section(section, *, first: bool = False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    if not first:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False

    for header in (section.header, section.even_page_header):
        header_p = header.paragraphs[0]
        header_p.clear()
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.paragraph_format.space_after = Pt(0)
        header_run = header_p.add_run("Kamilya LMS | индивидуальный B2B-договор")
        set_run_font(header_run, size=8.5, color=MUTED)

    for footer in (section.footer, section.even_page_footer):
        footer_p = footer.paragraphs[0]
        footer_p.clear()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(0)
        footer_run = footer_p.add_run("Проект до заполнения реквизитов  |  стр. ")
        set_run_font(footer_run, size=8.5, color=MUTED)
        add_page_field(footer_p)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_markdown_table(doc: Document, lines: list[str]):
    rows = parse_table(lines)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_index, row in enumerate(rows):
        for c_index, text in enumerate(row):
            paragraph = table.cell(r_index, c_index).paragraphs[0]
            paragraph.clear()
            add_inline_runs(paragraph, text, size=10)
    style_table(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_title(doc: Document, text: str, *, annex: bool):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4 if annex else 8)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper() if not annex else text)
    set_run_font(run, size=16 if not annex else 15, bold=True, color=NAVY)

    if not annex:
        note = doc.add_table(rows=1, cols=1)
        note.cell(0, 0).text = (
            "Проект для индивидуального согласования. Не является публичной офертой "
            "или договором с потребителем. Заполняемые поля выделены цветом."
        )
        set_table_geometry(note, [TABLE_WIDTH_DXA])
        note.style = "Table Grid"
        cell = note.cell(0, 0)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT_FILL)
        cell._tc.get_or_add_tcPr().append(shd)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for item in p.runs:
                set_run_font(item, size=9.5, bold=True, color=INK)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_markdown(doc: Document, source: Path, *, annex: bool):
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    first_heading = True
    while index < len(lines):
        raw = lines[index].rstrip()
        if not raw:
            index += 1
            continue
        if raw.startswith("| ") or raw.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if raw.startswith("# "):
            if first_heading:
                add_title(doc, raw[2:].strip(), annex=annex)
                first_heading = False
            index += 1
            continue
        if raw.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline_runs(paragraph, raw[3:].strip(), size=14, color=NAVY)
            index += 1
            continue
        if raw.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline_runs(paragraph, raw[4:].strip(), size=12, color=NAVY)
            index += 1
            continue
        if raw.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, raw[2:].strip())
            index += 1
            continue
        if re.match(r"^\d+\)\s", raw):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_runs(paragraph, re.sub(r"^\d+\)\s*", "", raw))
            index += 1
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.widow_control = True
        add_inline_runs(paragraph, raw)
        index += 1


def build():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    configure_section(doc.sections[0], first=True)
    doc.core_properties.title = "Договор об оказании услуг по доступу к Kamilya LMS"
    doc.core_properties.subject = "Индивидуальный B2B-договор и приложения"
    doc.core_properties.creator = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = "Рабочий проект до заполнения реквизитов и подписания"

    add_markdown(doc, SOURCES[0], annex=False)
    for source in SOURCES[1:]:
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        configure_section(section)
        add_markdown(doc, source, annex=True)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
