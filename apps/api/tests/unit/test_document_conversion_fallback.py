from __future__ import annotations

import hashlib

import pytest
from docx import Document
from pypdf import PdfWriter

from app.modules.ai.ingestion import (
    DocumentIngestion,
    DocumentOCRRequiredError,
    _local_convert,
)


@pytest.mark.asyncio
async def test_ingestion_preserves_canonical_document_source_revision(tmp_path) -> None:
    source = tmp_path / "policy.docx"
    source.write_bytes(b"canonical source blob")
    source_revision = f"document:{hashlib.sha256(source.read_bytes()).hexdigest()}"
    captured_chunks: list[dict] = []

    class ConverterStub:
        async def convert(self, file_path: str) -> dict:
            return {"markdown": "# Converted policy", "metadata": {"engine": "test"}}

    class ChunkerStub:
        def chunk_markdown(self, markdown: str, doc_id: str, filename: str) -> list[dict]:
            return [{"text": markdown, "metadata": {}}]

    class EmbeddingBatchStub:
        def as_lists(self) -> list[list[float]]:
            return [[0.1, 0.2]]

    class EmbeddingsStub:
        async def embed_documents_with_provenance(self, texts: list[str]):
            return EmbeddingBatchStub()

    class StoreStub:
        async def add_chunks(self, chunks, embedding_batch, *, tenant_id: str) -> int:
            captured_chunks.extend(chunks)
            return 0

    class SummarizerStub:
        async def summarize(self, markdown: str, doc_id: str, filename: str) -> dict:
            return {"summary": "test"}

    ingestion = DocumentIngestion(summaries_dir=str(tmp_path / "summaries"))
    ingestion.converter = ConverterStub()
    ingestion.chunker = ChunkerStub()
    ingestion.embeddings = EmbeddingsStub()
    ingestion.store = StoreStub()
    ingestion.summarizer = SummarizerStub()

    result = await ingestion.ingest_file(
        str(source),
        doc_id="document-id",
        tenant_id="00000000-0000-0000-0000-000000000001",
        source_revision=source_revision,
    )

    assert result["embeddings_written"] == 1
    assert captured_chunks[0]["metadata"]["source_revision"] == source_revision


@pytest.mark.asyncio
async def test_binary_document_is_never_indexed_as_placeholder(tmp_path) -> None:
    source = tmp_path / "policy.pdf"
    source.write_bytes(b"%PDF-1.7")

    with pytest.raises(RuntimeError, match="conversion is unavailable"):
        await _local_convert(str(source))


@pytest.mark.asyncio
async def test_plain_text_document_keeps_local_fallback(tmp_path) -> None:
    source = tmp_path / "policy.md"
    source.write_text("# Approved policy", encoding="utf-8")

    converted = await _local_convert(str(source))

    assert converted["markdown"] == "# Approved policy"
    assert converted["metadata"] == {
        "filename": "policy.md",
        "size": len("# Approved policy"),
        "pages": 0,
        "tables": 0,
        "engine": "plain_text",
        "engine_version": None,
        "fallback_used": False,
        "warnings": [],
    }


@pytest.mark.asyncio
async def test_docx_document_has_grounded_local_fallback(tmp_path) -> None:
    source = tmp_path / "policy.docx"
    document = Document()
    document.add_heading("Политика обучения", level=1)
    document.add_paragraph("Сотрудник должен набрать не менее 80 процентов.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Срок"
    table.cell(1, 1).text = "30 минут"
    document.save(source)

    converted = await _local_convert(str(source))

    assert "# Политика обучения" in converted["markdown"]
    assert "Сотрудник должен набрать не менее 80 процентов." in converted["markdown"]
    assert "| Параметр | Значение |" in converted["markdown"]
    assert "| Срок | 30 минут |" in converted["markdown"]
    assert converted["metadata"]["engine"] == "python-docx"
    assert converted["metadata"]["fallback_used"] is True


@pytest.mark.asyncio
async def test_pdf_document_has_local_text_fallback(tmp_path) -> None:
    source = tmp_path / "policy.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with source.open("wb") as target:
        writer.write(target)

    converted = await _local_convert(str(source))

    assert converted["markdown"] == ""
    assert converted["metadata"]["pages"] == 1
    assert converted["metadata"]["engine"] == "pypdf"
    assert converted["metadata"]["fallback_used"] is True


@pytest.mark.asyncio
async def test_scanned_pdf_requires_ocr_before_embedding(tmp_path) -> None:
    source = tmp_path / "scanned-policy.pdf"
    source.write_bytes(b"synthetic scanned fixture")

    class ScannedPDFConverter:
        async def convert(self, file_path: str) -> dict:
            return {
                "markdown": "",
                "metadata": {
                    "engine": "pypdf",
                    "fallback_used": True,
                    "pages": 2,
                },
            }

    class EmbeddingsMustNotRun:
        async def embed_documents_with_provenance(self, texts):
            raise AssertionError("empty embedding batch must not be sent")

    ingestion = DocumentIngestion()
    ingestion.converter = ScannedPDFConverter()
    ingestion.embeddings = EmbeddingsMustNotRun()

    with pytest.raises(DocumentOCRRequiredError, match="requires OCR"):
        await ingestion.ingest_file(
            str(source),
            doc_id="scanned-policy",
            tenant_id="00000000-0000-0000-0000-000000000001",
        )
